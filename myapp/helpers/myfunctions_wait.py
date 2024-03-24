"""This contains functions for specific tasks"""
import customtkinter as ctk
from tkinter import messagebox
from PIL import Image, ImageTk
import time
from appclasses.file_class import FileHandler
from threading import Thread
import re
import os
import validators
from docx import Document
import json
import img2pdf
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from fpdf import FPDF
import webbrowser
import tempfile
import socket
import psutil
import serial

def notify_user(msg=None):
    """ Pops a message when there is need to. Message to be displayed can be given as an arg """
    if msg is None:
        messagebox.showinfo("Alert!", f"I am a message button.")
    else:
        messagebox.showinfo("Alert!", "{}".format(msg))


def image_to_pdf_open(image_path):
    try:
        # Convert image to PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_file, open(image_path, "rb") as image_file:
            pdf_file.write(img2pdf.convert(image_file))
            pdf_path = pdf_file.name

        # Open the PDF in the default system browser
        webbrowser.open(f"file://{os.path.realpath(pdf_path)}")
    finally:
        # Delete the temporary PDF file after it has been opened
        if os.path.exists(pdf_path):
            os.remove(pdf_path)


def image_to_pdf_with_reportlab_open(image_path):
    pdf_path = "temp_output.pdf"  # Temporary PDF path

    # Create PDF using ReportLab
    c = canvas.Canvas(pdf_path, pagesize=letter)
    c.drawImage(image_path, 100, 100)  # Adjust coordinates as needed
    c.save()

    # Open the PDF in the default system browser
    webbrowser.open(f"file://{os.path.realpath(pdf_path)}")

    # Delete the temporary PDF file after it has been opened
    if os.path.exists(pdf_path):
        os.remove(pdf_path)


def image_to_pdf_with_pyfpdf_open(image_path):
    pdf_path = "temp_output.pdf"  # Temporary PDF path

    # Create PDF using PyFPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.image(image_path, x=10, y=10, w=200)  # Adjust coordinates as needed
    pdf.output(pdf_path)

    # Open the PDF in the default system browser
    webbrowser.open(f"file://{os.path.realpath(pdf_path)}")

    # Delete the temporary PDF file after it has been opened
    time.sleep(2)
    if os.path.exists(pdf_path):
        os.remove(pdf_path)


def create_image_obj(imagepath, size):
    """ This creates a CTkImage object that could be used on any of the app widgets
        It takes two args: imagepath (string)  to image and size (tuple) representing the width and height of image (w,h)
        Returns the image obj
    """
    image = Image.open(imagepath)
    img = ctk.CTkImage(image, size=size)
    # resized_image = image.resize(size, Image.ANTIALIAS)
    # img = ImageTk.PhotoImage(resized_image)
    return img


def save_files_to_app(filepath, location):
    """saves files to the user's app directory"""
    handler = FileHandler(filepath)
    response = handler.save_file(location)
    if response:
        notify_user('file saved successfully!')
    else:
        notify_user('Operation was not successful.')

    return


def is_valid_url(url):
    """checks for valid url while allowing localhost or ip urls"""
    is_valid = validators.url(url)
    if not is_valid:
        # Check if it contains 'localhost' or an IP address format
        ip_or_localhost_pattern = re.compile(
            r'^(https?|ftp)://([^\s/:]+|\[[^\s/]+\]|localhost)(:[0-9]+)?(\/|/([\w#!:.?+=&%@!\-/]))?$')
        if ip_or_localhost_pattern.match(url):
            return True
        return False
    return True


def addspace(obj_position, obj_size):
    """Gets the position of the object and then adds up the size of the object to
     give the position of the point where the span of the obj ended
    """
    return obj_position + obj_size


def get_massxxx(scale_settings=None):
    # ser = serial.Serial(port='/dev/ttyS0', ## if using Linux serial
    if scale_settings:
        port = scale_settings['port']
        baudrate = scale_settings['baudrate']
        parity = scale_settings['parity']
        stopbits = scale_settings['stopbits']
        bytesize = scale_settings['bytesize']
    else:
        port = "COM2"
        baudrate = 9600
        parity = serial.PARITY_NONE
        stopbits = serial.STOPBITS_ONE
        bytesize = serial.EIGHTBITS

    ser = serial.Serial(port,
                        baudrate=baudrate,
                        parity=parity,
                        stopbits=stopbits,
                        bytesize=bytesize
                        )

    if not ser.isOpen():
        ser.open()

    ser.write(b'\nSI\n')
    time.sleep(1)  # give moment for balance to settle

    value = ser.read(ser.in_waiting)

    value = decode_byte(value)

    value = digits_between_sequences(value, str(32))

    if ser.isOpen():
        ser.close()

    return int(value)


def is_ngrok_running():
    for proc in psutil.process_iter(['pid', 'name']):
        if 'ngrok' in proc.info['name']:
            return True
    return False


def get_mass(scale_settings=None):
    result = {'status': 2, 'data': None, 'message': None}

    try:
        if scale_settings:
            port = scale_settings.get('port', "COM2")
            baudrate = scale_settings.get('baudrate', 9600)
            parity = scale_settings.get('parity', 'N')
            stopbits = scale_settings.get('stopbits', 1)
            bytesize = scale_settings.get('bytesize', 8)
        else:
            port = "COM2"
            baudrate = 9600
            parity = 'N'
            stopbits = 1
            bytesize = 8

        ser = serial.Serial(port,
                            baudrate=baudrate,
                            parity=parity,
                            stopbits=stopbits,
                            bytesize=bytesize
                            )

        if not ser.isOpen():
            ser.open()

        ser.write(b'\nSI\n')
        time.sleep(1)  # give moment for balance to settle

        value = ser.read(ser.in_waiting)

        value = decode_byte(value)

        value = digits_between_sequences(value, str(32))

        result['status'] = 1
        result['data'] = int(value)
        result['message'] = "Success"

    except Exception as e:
        result['message'] = str(e)

    finally:
        if 'ser' in locals():
            if ser.isOpen():
                ser.close()

    return result


def decode_byte(byte_value):
    """Decodes the bytes data gotten from the scale """

    x = str(byte_value).split("'")
    x = x[1].split("\\")
    new_x = x[5:-3]
    k = ''
    for w in new_x:
        k += w[2:]
    return k


def digits_between_sequences(larger_sequence, sequence_to_find):
    """
    Finds digits between consecutive occurrences of a given sequence within a larger sequence.

    Args:
    - larger_sequence (str): The larger sequence to search within.
    - sequence_to_find (str): The specific sequence to find within the larger sequence.

    Returns:
    - str: A string containing the digits found between consecutive occurrences
           of the specified sequence within the larger sequence.
    """
    result = ""
    start = 0
    while True:
        # Find the next occurrence of the sequence within the larger sequence
        start_idx = larger_sequence.find(sequence_to_find, start)
        if start_idx == -1:  # If no more occurrences found, break the loop
            break

        # Move the starting index after the current sequence to look for the next occurrence
        start = start_idx + len(sequence_to_find)

        # Find the end index for the next occurrence of the sequence
        end_idx = larger_sequence.find(sequence_to_find, start)
        if end_idx == -1:  # If no more occurrences found, break the loop
            break

        # Extract the digits between the two occurrences and add them to the result
        result += larger_sequence[start:end_idx]

    return result[:7]


def is_internet_connected():
    """checks if the system is connected to the internet or not"""
    try:
        # Attempt to connect to a well-known website (in this case, Google's DNS server)
        socket.create_connection(("8.8.8.8", 53), timeout=3)
        return True  # If connection succeeds, return True
    except OSError:
        pass
    return False  # If connection fails, return False


def is_valid_email(email: str) -> bool:
    """
    Checks if the given string is a valid email address.
    Returns True if it's a valid email, False otherwise.
    """
    # Regular expression pattern for validating an email address
    email_pattern = r"^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$"

    # Using the re.match function to check if the email matches the pattern
    if re.match(email_pattern, email):
        return True
    else:
        return False


def thread_request(func, *args, **kwargs):
    """Starts a thread that invokes functions for specific actions"""
    # Start a thread to handle the database operation
    thread = Thread(target=func, args=args, kwargs=kwargs)
    thread.daemon = True  # Daemonize the thread to avoid issues on application exit
    thread.start()


def get_ip_address():
    """
    Retrieves the local IP address of the system by creating a socket connection
    to a remote server (in this case, Google's DNS server) and extracting the
    local IP address connected to it.

    Returns:
        str: The IP address of the system if retrieved successfully.
             If unable to retrieve the IP address, returns a message indicating
             the failure.
    """
    try:
        # Create a socket object
        sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        # Connect to a remote server (does not send any packets)
        sock.connect(("8.8.8.8", 80))
        # Get the local IP address connected to the remote host
        ip_address = sock.getsockname()[0]
        # Close the socket
        sock.close()
        return ip_address
    except socket.error:
        return "Unable to retrieve IP address"


def store_app_settings(settings):
    file_path = "app_settings.json"

    try:
        with open(file_path, 'w') as file:
            json.dump(settings, file, indent=4)

        message = f"Settings saved to '{file_path}'"
        status = 1
    except Exception as e:
        message = f"Error writing settings to '{file_path}': {e}"
        status = 2

    return {'status': status, 'message': message, 'data': None}


def read_app_settings():
    file_path = "app_settings.json"
    try:
        with open(file_path, 'r') as file:
            settings = json.load(file)
        data = settings
        status = 1
        message = "success"
    except FileNotFoundError:
        data = None
        status = 2
        message = f"File '{file_path}' not found."
    except json.JSONDecodeError:
        data = None
        status = 2
        message = f"Error decoding JSON from '{file_path}'. Check if the file contains valid JSON."

    return {'status': status, 'message': message, 'data': data}


def update_template(replacements, file_path, output_path):
    doc = Document(file_path)

    # replace variables in the doc file body
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, str(value))
                        inline[i].text = text

    # replace variables in the table cells if any
    if doc.tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, str(value))
    # fill the table cells with product data
    if doc.tables:
        for table in doc.tables:
            if len(replacements['products']) > 0:
                products = replacements['products']
                for i, row_data in enumerate(products):
                    if i < len(table.rows):
                        headers = list(row_data.keys())
                        headers.insert(0, 'S/N')
                        for j, header in enumerate(headers):
                            if j == 0:
                                table.rows[i + 1].cells[j].text = f'{j+1}'
                            else:
                                table.rows[i+1].cells[j].text = f'{row_data[header]}'
                break

    doc.save(output_path)
    return


